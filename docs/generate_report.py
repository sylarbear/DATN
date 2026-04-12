# -*- coding: utf-8 -*-
"""
Generate full graduation thesis report (Báo cáo thực tập tốt nghiệp)
Based on template: Mau-bao-cao-chuyen-de-thuc-tap_New.docx
"""
import os
import copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DOCS_DIR = r"g:\xamp\htdocs\DATN\docs"
TEMPLATE = os.path.join(DOCS_DIR, "mau_bao_cao", "Mau-bao-cao-chuyen-de-thuc-tap_New.docx")
SCREENSHOTS_DIR = os.path.join(DOCS_DIR, "screenshots")
UML_IMAGES_DIR = os.path.join(DOCS_DIR, "uml_images")
OUTPUT = os.path.join(DOCS_DIR, "bao_cao_thuc_tap.docx")

# Load template to get styles
doc = Document(TEMPLATE)

# ============================================================
# HELPER: Clear all content from template
# ============================================================
body = doc.element.body
# Remove all paragraphs and tables
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag in ('p', 'tbl', 'sectPr'):
        if tag == 'sectPr':
            continue  # keep section properties
        body.remove(child)

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def add_para(text, style='Normal', bold=False, italic=False, align=None, size=None, space_after=None):
    """Add a paragraph with style"""
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def add_heading1(text):
    """Chapter heading - Heading 1"""
    return add_para(text, 'Heading 1')

def add_level1(text):
    """Section heading - _level 1"""
    return add_para(text, '_level 1')

def add_level2(text):
    """Sub-section heading - _level2"""
    return add_para(text, '_level2')

def add_level3(text):
    """Sub-sub-section heading - _level3"""
    return add_para(text, '_level3')

def add_body(text):
    """Body text - _doanVB"""
    return add_para(text, '_doanVB')

def add_bullet(text):
    """Bullet point"""
    return add_para(text, 'List Paragraph')

def add_image(img_path, caption=None, width=5.5):
    """Add image with caption"""
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = add_para(caption, 'Normal', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    else:
        add_para(f"[Hình: {os.path.basename(img_path)}]", 'Normal', italic=True)

def add_page_break():
    doc.add_page_break()

# ============================================================
# TRANG BÌA
# ============================================================
add_para("BỘ CÔNG THƯƠNG", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
add_para("TRƯỜNG CAO ĐẲNG CÔNG THƯƠNG TPHCM", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("KHOA CÔNG NGHỆ THÔNG TIN", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)

for _ in range(3):
    add_para("", 'Normal')

add_para("BÁO CÁO THỰC TẬP", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22)
add_para("TỐT NGHIỆP", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=22)

add_para("", 'Normal')
add_para("ĐỀ TÀI: XÂY DỰNG WEBSITE HỌC TIẾNG ANH TRỰC TUYẾN", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para("ENGLISH LEARNING", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)

for _ in range(3):
    add_para("", 'Normal')

add_para("GV hướng dẫn: Vũ Thị Hường", 'Normal', size=14)
add_para("Sinh viên thực hiện: Phan Quang Thuật", 'Normal', size=14)
add_para("Mã số sinh viên: 2120110351", 'Normal', size=14)
add_para("Lớp: CCQ2011E", 'Normal', size=14)
add_para("Khoá: K44", 'Normal', size=14)

for _ in range(4):
    add_para("", 'Normal')

add_para("TPHCM, tháng 4 năm 2026", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_page_break()

# ============================================================
# LỜI CẢM ƠN
# ============================================================
add_heading1("LỜI CẢM ƠN")
add_body("Trong quá trình thực hiện đề tài \"Xây dựng website học tiếng Anh trực tuyến English Learning\", em đã nhận được rất nhiều sự giúp đỡ và hướng dẫn tận tình.")
add_body("Trước hết, em xin gửi lời cảm ơn chân thành đến cô Vũ Thị Hường – Giảng viên hướng dẫn, người đã tận tình chỉ bảo, định hướng và đóng góp những ý kiến quý báu trong suốt quá trình em thực hiện đề tài. Những góp ý của cô đã giúp em hoàn thiện sản phẩm một cách tốt nhất.")
add_body("Em xin cảm ơn quý Thầy Cô trong Khoa Công nghệ Thông tin – Trường Cao đẳng Công Thương TPHCM đã truyền đạt cho em những kiến thức nền tảng vững chắc trong suốt thời gian học tập tại trường, tạo điều kiện cho em có thể áp dụng vào thực tế.")
add_body("Cuối cùng, em xin gửi lời cảm ơn đến gia đình, bạn bè đã luôn động viên, hỗ trợ em trong suốt quá trình học tập và thực hiện đề tài.")
add_body("Trong quá trình thực hiện, do kiến thức và kinh nghiệm còn hạn chế nên đề tài không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của quý Thầy Cô để em có thể hoàn thiện hơn.")
add_para("", 'Normal')
add_para("TPHCM, tháng 4 năm 2026", 'Normal', italic=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("Sinh viên thực hiện", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("Phan Quang Thuật", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_page_break()

# ============================================================
# NHẬN XÉT GVHD (để trống cho cô điền)
# ============================================================
add_para("NHẬN XÉT KẾT QUẢ THỰC TẬP", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
add_para("(Dành cho Giảng viên hướng dẫn sinh viên thực tập)", 'Normal', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("", 'Normal')
add_bullet("Họ tên sinh viên thực tập: Phan Quang Thuật")
add_bullet("Mã số sinh viên: 2120110351")
add_bullet("Lớp: CCQ2011E – Khóa: K44")
add_bullet("Giảng viên hướng dẫn thực tập: Vũ Thị Hường")
add_bullet("Sau thời gian sinh viên Phan Quang Thuật thực tập tốt nghiệp, tôi có nhận xét sau:")
add_para("", 'Normal')
add_para("Về ý thức, thái độ của sinh viên:", 'Normal', bold=True)
for _ in range(3):
    add_para(".......................................................................................................................................", 'Normal')
add_para("Về năng lực chuyên môn:", 'Normal', bold=True)
for _ in range(3):
    add_para(".......................................................................................................................................", 'Normal')
add_para("Kết luận:", 'Normal', bold=True)
add_para(".......................................................................................................................................", 'Normal')
add_para("Nhận xét:", 'Normal', bold=True)
add_para(".......................................................................................................................................", 'Normal')
add_para("Điểm:", 'Normal', bold=True)
add_para(".......................................................................................................................................", 'Normal')
add_para("", 'Normal')
add_para("TPHCM, Ngày…… tháng…… năm 2026", 'Normal', align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("Giảng viên hướng dẫn", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("", 'Normal')
add_para("", 'Normal')
add_para("Vũ Thị Hường", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_page_break()

# ============================================================
# ĐỀ CƯƠNG THỰC TẬP
# ============================================================
add_heading1("ĐỀ CƯƠNG THỰC TẬP")

add_para("1. Nhiệm vụ / Đề tài:", 'Normal', bold=True)
add_body("XÂY DỰNG WEBSITE HỌC TIẾNG ANH TRỰC TUYẾN ENGLISH LEARNING")

add_para("2. Mục tiêu công việc / đề tài:", 'Normal', bold=True)
add_bullet("Đề tài thuộc lĩnh vực phát triển ứng dụng web, giải quyết bài toán hỗ trợ người dùng học tiếng Anh trực tuyến thông qua các phương pháp học tương tác: từ vựng, ngữ pháp, luyện nghe, đọc, nói.")
add_bullet("Ứng dụng mang ý nghĩa thiết thực trong giáo dục, giúp người học tiếp cận việc học tiếng Anh một cách dễ dàng, miễn phí hoặc với chi phí thấp, mọi lúc mọi nơi thông qua trình duyệt web.")
add_bullet("Nắm vững kiến thức về xây dựng website theo mô hình MVC sử dụng PHP thuần, MySQL, HTML/CSS/JavaScript.")
add_bullet("Xây dựng một trang web thân thiện, dễ sử dụng, tích hợp AI đánh giá kỹ năng nói cho người học tiếng Anh.")
add_bullet("Xây dựng hệ thống quản lý membership (ví điện tử, nạp tiền, mua gói) chuyên nghiệp và bảo mật.")

add_para("3. Đối tượng nghiên cứu:", 'Normal', bold=True)
add_bullet("Đối tượng nghiên cứu của đề tài là lập trình web nói chung và ngôn ngữ lập trình PHP, HTML, CSS, JavaScript nói riêng.")
add_bullet("Cơ sở dữ liệu MySQL, quản lý phiên bản với Git/GitHub.")
add_bullet("Tích hợp API: Google OAuth 2.0, OpenAI GPT, Web Speech API, VietQR.")
add_bullet("Nghiên cứu nhu cầu của người dùng học tiếng Anh online để xây dựng giao diện thân thiện, trải nghiệm tốt.")

add_para("4. Nội dung chính của thực tập:", 'Normal', bold=True)
add_para("Chương 1: Tổng quan về đề tài", 'Normal')
add_para("Chương 2: Lý thuyết áp dụng theo lĩnh vực thực tập", 'Normal')
add_bullet("Nghiên cứu tổng quan về đề tài.")
add_bullet("Nghiên cứu cơ sở lý thuyết của đề tài.")
add_para("Chương 3: Nội dung thực tập", 'Normal')
add_bullet("Phân tích và thiết kế hệ thống (UML).")
add_bullet("Xây dựng cơ sở dữ liệu.")
add_bullet("Xây dựng giao diện ứng dụng web.")
add_bullet("Xây dựng chức năng cho ứng dụng.")
add_bullet("Kiểm thử và kiểm tra, sửa lỗi.")
add_para("Chương 4: Kết luận và kiến nghị", 'Normal')
add_bullet("Kết quả thực hiện: Ưu điểm, Hạn chế.")
add_bullet("Hướng phát triển trong tương lai.")

add_para("5. Tiến độ thực hiện thực tập:", 'Normal', bold=True)
# Tiến độ table
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
headers = ['TT', 'Thời gian', 'Nội dung công việc', 'Kết quả dự kiến']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

schedule = [
    ['1', 'Tuần 1-2', 'Tìm hiểu yêu cầu, viết đề cương, thiết kế CSDL', 'Đề cương + ERD hoàn thành'],
    ['2', 'Tuần 3-5', 'Xây dựng giao diện và chức năng chính (Topics, Lessons, Tests, Grammar)', 'Hoàn thành module học tập'],
    ['3', 'Tuần 6-8', 'Tích hợp Speaking AI, Wallet, Membership, Support', 'Hoàn thành toàn bộ chức năng'],
    ['4', 'Tuần 9-10', 'Kiểm thử, sửa lỗi, viết tài liệu UML, chụp wireframe', 'Sản phẩm hoàn chỉnh + báo cáo'],
    ['5', 'Tuần 11-12', 'Hoàn thiện báo cáo, nộp bài, chuẩn bị bảo vệ', 'Báo cáo hoàn chỉnh'],
]
for r, data in enumerate(schedule):
    for c, text in enumerate(data):
        table.rows[r+1].cells[c].text = text

add_page_break()

# ============================================================
# MỤC LỤC (placeholder - user tạo bằng Word)
# ============================================================
add_para("MỤC LỤC", 'Normal', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
add_para("", 'Normal')
add_para("(Nhấn chuột phải → Update Field để cập nhật mục lục trong Word)", 'Normal', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_page_break()

# ============================================================
# CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI
# ============================================================
add_heading1("TỔNG QUAN VỀ ĐỀ TÀI")

add_level1("1.1. Giới thiệu đề tài")
add_body("Trong bối cảnh toàn cầu hóa và hội nhập quốc tế ngày càng sâu rộng, tiếng Anh đã trở thành ngôn ngữ giao tiếp quốc tế không thể thiếu. Tại Việt Nam, nhu cầu học tiếng Anh ngày càng tăng cao, đặc biệt ở giới trẻ và sinh viên. Tuy nhiên, việc tiếp cận các khóa học tiếng Anh chất lượng với chi phí hợp lý vẫn còn là thách thức đối với nhiều người.")
add_body("Sự phát triển nhanh chóng của công nghệ thông tin và Internet đã mở ra cơ hội xây dựng các nền tảng học tập trực tuyến (e-learning), cho phép người học tiếp cận kiến thức mọi lúc, mọi nơi. Các ứng dụng học tiếng Anh trực tuyến giúp người dùng chủ động trong việc học, tiết kiệm thời gian và chi phí so với phương pháp truyền thống.")
add_body("Đề tài \"Xây dựng website học tiếng Anh trực tuyến English Learning\" được thực hiện nhằm phát triển một nền tảng web toàn diện, hỗ trợ người dùng học tiếng Anh qua nhiều kỹ năng: từ vựng, ngữ pháp, nghe, đọc, nói. Đặc biệt, hệ thống tích hợp trí tuệ nhân tạo (AI) để đánh giá kỹ năng nói, tạo trải nghiệm học tập tương tác và hiệu quả hơn.")

add_level1("1.2. Lý do chọn đề tài")
add_body("Việc chọn đề tài này xuất phát từ những lý do sau:")
add_bullet("Nhu cầu thực tế: Thị trường học tiếng Anh online tại Việt Nam đang tăng trưởng mạnh mẽ. Các nền tảng như Duolingo, ELSA Speak đã chứng minh tính khả thi và hiệu quả của việc học ngoại ngữ qua ứng dụng công nghệ.")
add_bullet("Ứng dụng công nghệ AI: Sự phát triển của AI (đặc biệt là GPT và Web Speech API) cho phép đánh giá kỹ năng nói tự động, mở ra hướng tiếp cận mới trong giáo dục ngoại ngữ.")
add_bullet("Tích hợp đa kỹ năng: Khác với các ứng dụng chỉ tập trung một kỹ năng, English Learning tích hợp tất cả 4 kỹ năng (Nghe, Nói, Đọc, Viết) và thêm từ vựng + ngữ pháp trong một nền tảng duy nhất.")
add_bullet("Hệ thống quản lý tài chính: Tích hợp ví điện tử, nạp tiền qua QR code, mua gói membership – cung cấp mô hình kinh doanh bền vững cho nền tảng giáo dục.")
add_bullet("Phù hợp với chuyên ngành: Đề tài cho phép áp dụng toàn diện các kiến thức đã học: lập trình web (PHP, HTML, CSS, JS), cơ sở dữ liệu (MySQL), mô hình MVC, và tích hợp API.")

add_level1("1.3. Mục tiêu đề tài")
add_level2("1.3.1. Mục tiêu tổng quát")
add_body("Xây dựng một website học tiếng Anh trực tuyến hoàn chỉnh, có giao diện thân thiện, nhiều tính năng học tập tương tác, và hệ thống quản trị + quản lý tài chính chuyên nghiệp.")

add_level2("1.3.2. Mục tiêu cụ thể")
add_bullet("Xây dựng hệ thống quản lý nội dung học tập: chủ đề từ vựng, bài học, bài kiểm tra, ngữ pháp.")
add_bullet("Phát triển tính năng Flashcard giúp người dùng học từ vựng hiệu quả.")
add_bullet("Tích hợp AI đánh giá kỹ năng nói (Speaking) với điểm số chi tiết: Pronunciation, Fluency, Accuracy.")
add_bullet("Xây dựng hệ thống bài kiểm tra đa dạng: Quiz, Listening, Reading với đếm thời gian.")
add_bullet("Phát triển hệ thống ví điện tử: nạp tiền, mua gói Pro, rút tiền, hoàn tiền.")
add_bullet("Xây dựng hệ thống membership (Free/Pro) với mã kích hoạt và thanh toán qua ví.")
add_bullet("Phát triển dashboard thống kê tiến độ học tập, bảng xếp hạng, hệ thống XP + Level.")
add_bullet("Xây dựng trang quản trị (Admin) với đầy đủ chức năng CRUD và quản lý tài chính.")

add_level1("1.4. Phạm vi đề tài")
add_body("Đề tài tập trung vào việc xây dựng và phát triển ứng dụng web với các phạm vi sau:")
add_bullet("Nền tảng: Website chạy trên trình duyệt (responsive, tương thích PC và mobile).")
add_bullet("Ngôn ngữ & công nghệ: PHP (thuần MVC), MySQL, HTML5, CSS3, JavaScript, AJAX.")
add_bullet("Tích hợp bên thứ ba: Google OAuth 2.0, OpenAI GPT API, Web Speech API, VietQR.")
add_bullet("Đối tượng sử dụng: Học sinh, sinh viên, người đi làm muốn học tiếng Anh trực tuyến.")
add_bullet("Phạm vi chức năng: Học từ vựng, ngữ pháp, luyện nói, làm bài kiểm tra, quản lý ví, mua gói membership.")
add_page_break()

# ============================================================
# CHƯƠNG 2: LÝ THUYẾT ÁP DỤNG
# ============================================================
add_heading1("LÝ THUYẾT ÁP DỤNG THEO LĨNH VỰC THỰC TẬP")

add_level1("2.1. Tổng quan về đề tài")

add_level2("2.1.1. Tổng quan về E-Learning")
add_body("E-Learning (Electronic Learning) là hình thức đào tạo trực tuyến sử dụng công nghệ thông tin và Internet để truyền tải nội dung học tập. E-Learning cho phép người học tiếp cận bài giảng, tài liệu, bài kiểm tra mọi lúc mọi nơi thông qua các thiết bị kết nối Internet.")
add_body("Ưu điểm của E-Learning bao gồm: linh hoạt về thời gian và địa điểm học, chi phí thấp hơn so với đào tạo truyền thống, khả năng cá nhân hóa lộ trình học tập, và tích hợp công nghệ AI để nâng cao trải nghiệm.")
add_body("Tại Việt Nam, thị trường E-Learning đang phát triển mạnh với các nền tảng như VioEdu, Topica, Moon.vn. Riêng lĩnh vực học tiếng Anh online, các ứng dụng như Duolingo, ELSA Speak, Cake đã thu hút hàng triệu người dùng Việt Nam.")

add_level2("2.1.2. Tổng quan về ứng dụng học tiếng Anh trực tuyến")
add_body("Các ứng dụng học tiếng Anh trực tuyến hiện nay thường tập trung vào các phương pháp sau:")
add_bullet("Gamification: Áp dụng các yếu tố trò chơi (điểm XP, streak, level, leaderboard) để tạo động lực học tập.")
add_bullet("Spaced Repetition: Lặp lại ngắt quãng – ôn tập từ vựng theo chu kỳ tối ưu để ghi nhớ lâu dài.")
add_bullet("AI-powered Assessment: Sử dụng AI để đánh giá phát âm, ngữ pháp, và cung cấp phản hồi tức thì.")
add_bullet("Multi-skill Integration: Tích hợp đa kỹ năng (Nghe, Nói, Đọc, Viết) trong một nền tảng duy nhất.")
add_body("English Learning áp dụng tất cả các phương pháp trên, kết hợp với hệ thống quản lý membership và ví điện tử, tạo ra một nền tảng học tập toàn diện.")

add_level1("2.2. Cơ sở lý thuyết")

add_level2("2.2.1. Ngôn ngữ lập trình PHP")
add_body("PHP (PHP: Hypertext Preprocessor) là ngôn ngữ lập trình kịch bản phía máy chủ (server-side scripting language) được sử dụng rộng rãi trong phát triển web. PHP được thiết kế để tạo ra các trang web động, xử lý form, quản lý session, cookie, và tương tác với cơ sở dữ liệu.")
add_body("Đặc điểm nổi bật của PHP:")
add_bullet("Mã nguồn mở, miễn phí, cộng đồng hỗ trợ lớn.")
add_bullet("Hỗ trợ nhiều hệ quản trị CSDL: MySQL, PostgreSQL, SQLite, Oracle.")
add_bullet("Tích hợp tốt với các web server: Apache, Nginx, IIS.")
add_bullet("Cú pháp đơn giản, dễ học, dễ triển khai.")
add_body("Trong đề tài này, PHP được sử dụng phiên bản 8.2 trở lên, áp dụng mô hình MVC (Model-View-Controller) để tổ chức mã nguồn một cách khoa học và dễ bảo trì.")

add_level2("2.2.2. Mô hình MVC (Model-View-Controller)")
add_body("MVC là một design pattern phổ biến trong phát triển web, chia ứng dụng thành 3 thành phần chính:")
add_bullet("Model: Xử lý logic nghiệp vụ và tương tác với cơ sở dữ liệu. Ví dụ: User Model xử lý đăng ký, đăng nhập; Topic Model quản lý chủ đề học.")
add_bullet("View: Hiển thị giao diện người dùng (HTML, CSS, JavaScript). Nhận dữ liệu từ Controller và render ra trang web.")
add_bullet("Controller: Tiếp nhận request từ người dùng, gọi Model xử lý, và chuyển kết quả đến View để hiển thị.")
add_body("Ưu điểm của MVC: tách biệt logic nghiệp vụ và giao diện, dễ bảo trì và mở rộng, hỗ trợ làm việc nhóm hiệu quả, tái sử dụng code.")

add_level2("2.2.3. Hệ quản trị CSDL MySQL")
add_body("MySQL là hệ quản trị cơ sở dữ liệu quan hệ (RDBMS) mã nguồn mở phổ biến nhất thế giới. MySQL sử dụng ngôn ngữ SQL (Structured Query Language) để quản lý và truy vấn dữ liệu.")
add_body("Đặc điểm của MySQL:")
add_bullet("Hiệu suất cao, xử lý nhanh, hỗ trợ transaction (InnoDB engine).")
add_bullet("Hỗ trợ ACID (Atomicity, Consistency, Isolation, Durability) đảm bảo tính toàn vẹn dữ liệu.")
add_bullet("Bảo mật tốt với hệ thống phân quyền chi tiết.")
add_bullet("Tích hợp tốt với PHP thông qua PDO (PHP Data Objects).")
add_body("Trong đề tài, MySQL được sử dụng với charset utf8mb4 để hỗ trợ đầy đủ ký tự tiếng Việt và emoji. Database bao gồm 23 bảng quản lý toàn bộ dữ liệu hệ thống.")

add_level2("2.2.4. HTML5, CSS3 và JavaScript")
add_body("HTML5 (HyperText Markup Language 5) là phiên bản mới nhất của ngôn ngữ đánh dấu dùng để cấu trúc nội dung trang web. HTML5 bổ sung các thẻ semantic (header, nav, section, article, footer) và hỗ trợ multimedia (audio, video) không cần plugin.")
add_body("CSS3 (Cascading Style Sheets 3) dùng để định dạng giao diện web. CSS3 bổ sung nhiều tính năng mạnh mẽ: Flexbox, Grid Layout, Animation, Gradient, Media Queries (responsive design).")
add_body("JavaScript là ngôn ngữ lập trình phía client, cho phép tạo tương tác động trên trang web. Trong đề tài, JavaScript được sử dụng cho: xử lý AJAX (gọi API không reload trang), Web Speech API (ghi âm giọng nói), Chart.js (biểu đồ thống kê), và DOM manipulation.")

add_level2("2.2.5. Google OAuth 2.0")
add_body("OAuth 2.0 là giao thức ủy quyền (authorization) cho phép ứng dụng bên thứ ba truy cập tài nguyên của người dùng mà không cần biết mật khẩu. Google OAuth 2.0 cho phép người dùng đăng nhập bằng tài khoản Google, giúp:")
add_bullet("Đơn giản hóa quy trình đăng ký/đăng nhập cho người dùng.")
add_bullet("Tăng bảo mật (không lưu mật khẩu Google, chỉ lưu token).")
add_bullet("Lấy thông tin profile (tên, email, avatar) tự động.")

add_level2("2.2.6. OpenAI GPT API")
add_body("OpenAI GPT (Generative Pre-trained Transformer) là mô hình ngôn ngữ lớn (LLM) có khả năng hiểu và tạo văn bản tự nhiên. Trong đề tài, GPT API được sử dụng để:")
add_bullet("Đánh giá kỹ năng nói: Phân tích transcript từ Web Speech API, chấm điểm Pronunciation, Fluency, Accuracy.")
add_bullet("Cung cấp feedback chi tiết: Gợi ý cải thiện phát âm, ngữ pháp, và từ vựng.")
add_bullet("Chatbot hỗ trợ: Trả lời câu hỏi về tiếng Anh, giải thích ngữ pháp.")

add_level2("2.2.7. Web Speech API")
add_body("Web Speech API là API trình duyệt cho phép tích hợp tính năng nhận dạng giọng nói (Speech Recognition) và tổng hợp giọng nói (Speech Synthesis) trực tiếp trên trình duyệt web mà không cần cài đặt phần mềm bổ sung.")
add_body("Trong đề tài, Speech Recognition API được sử dụng để ghi nhận giọng nói của người dùng khi luyện nói tiếng Anh, chuyển đổi thành văn bản (transcript), sau đó gửi lên server để AI đánh giá.")

add_level2("2.2.8. VietQR – Thanh toán QR Code")
add_body("VietQR là tiêu chuẩn QR Code thanh toán được Ngân hàng Nhà nước Việt Nam ban hành. VietQR API cho phép tạo mã QR chuyển khoản tự động với đầy đủ thông tin: số tài khoản, ngân hàng, số tiền, nội dung chuyển khoản.")
add_body("Trong đề tài, VietQR được tích hợp vào tính năng nạp tiền vào ví, giúp người dùng chuyển khoản nhanh chóng bằng cách quét mã QR qua ứng dụng ngân hàng.")
add_page_break()

# ============================================================
# CHƯƠNG 3: NỘI DUNG THỰC TẬP
# ============================================================
add_heading1("NỘI DUNG THỰC TẬP")

add_level1("3.1. Phân tích và thiết kế hệ thống")

add_level2("3.1.1. Sơ đồ Use Case")
add_body("Sơ đồ Use Case mô tả các chức năng chính của hệ thống và mối quan hệ giữa các actor (tác nhân) với hệ thống. Hệ thống English Learning có 4 actor chính:")
add_bullet("Guest (Khách): Người dùng chưa đăng nhập – có thể xem chủ đề, đăng ký, đăng nhập.")
add_bullet("User (Học viên Free): Người dùng đã đăng nhập – có thể học từ vựng, làm quiz, flashcard, ngữ pháp, bookmark, quản lý ví.")
add_bullet("Pro User (Hội viên Pro): Kế thừa tất cả quyền User + truy cập Listening, Reading, Speaking.")
add_bullet("Admin (Quản trị viên): Kế thừa tất cả quyền Pro + quản lý hệ thống, duyệt đơn, quản lý ví.")

img = os.path.join(UML_IMAGES_DIR, "use_case.png")
add_image(img, "Hình 3.1: Sơ đồ Use Case tổng quan", 5.0)

add_level2("3.1.2. Sơ đồ Sequence")
add_body("Sơ đồ Sequence mô tả luồng tương tác giữa các thành phần trong hệ thống theo thời gian. Dưới đây là các luồng chính:")

add_level3("a) Luồng đăng nhập")
img = os.path.join(UML_IMAGES_DIR, "sequence_login.png")
add_image(img, "Hình 3.2: Sequence Diagram – Đăng nhập", 5.5)

add_level3("b) Luồng nạp tiền vào ví")
img = os.path.join(UML_IMAGES_DIR, "sequence_deposit.png")
add_image(img, "Hình 3.3: Sequence Diagram – Nạp tiền vào ví", 5.5)

add_level3("c) Luồng mua gói Pro bằng ví")
img = os.path.join(UML_IMAGES_DIR, "sequence_purchase.png")
add_image(img, "Hình 3.4: Sequence Diagram – Mua gói Pro", 5.5)

add_level3("d) Luồng hủy đơn và hoàn tiền")
img = os.path.join(UML_IMAGES_DIR, "sequence_cancel.png")
add_image(img, "Hình 3.5: Sequence Diagram – Hủy đơn + Hoàn tiền", 5.5)

add_level2("3.1.3. Sơ đồ State (Trạng thái)")
add_body("Sơ đồ State mô tả các trạng thái có thể có của một đối tượng và sự chuyển đổi giữa chúng.")

add_level3("a) Trạng thái Đơn hàng (Membership Order)")
img = os.path.join(UML_IMAGES_DIR, "state_order.png")
add_image(img, "Hình 3.6: State Diagram – Membership Order", 4.0)

add_level3("b) Trạng thái Giao dịch Ví (Wallet Transaction)")
img = os.path.join(UML_IMAGES_DIR, "state_wallet.png")
add_image(img, "Hình 3.7: State Diagram – Wallet Transaction", 4.0)

add_level3("c) Trạng thái Support Ticket")
img = os.path.join(UML_IMAGES_DIR, "state_ticket.png")
add_image(img, "Hình 3.8: State Diagram – Support Ticket", 4.0)

add_level2("3.1.4. Thiết kế cơ sở dữ liệu")
add_body("Hệ thống sử dụng MySQL với 23 bảng chính. Dưới đây là danh sách các bảng và mô tả:")

# Database tables
db_table = doc.add_table(rows=16, cols=3)
db_table.style = 'Table Grid'
db_headers = ['STT', 'Tên bảng', 'Mô tả']
for i, h in enumerate(db_headers):
    db_table.rows[0].cells[i].text = h
    for p in db_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
db_data = [
    ['1', 'users', 'Thông tin người dùng (username, email, balance, membership, XP...)'],
    ['2', 'topics', 'Chủ đề học (Travel, Food, Business...)'],
    ['3', 'vocabularies', 'Từ vựng (word, pronunciation, meaning, example)'],
    ['4', 'lessons / lesson_contents', 'Bài học + nội dung đa phương tiện'],
    ['5', 'tests / questions', 'Bài kiểm tra + câu hỏi (quiz, listening, reading)'],
    ['6', 'grammar_lessons / grammar_questions', 'Bài ngữ pháp + câu hỏi quiz'],
    ['7', 'speaking_prompts / speaking_attempts', 'Đề luyện nói + kết quả AI chấm điểm'],
    ['8', 'membership_plans', 'Gói Pro (1 tháng, 3 tháng, 12 tháng)'],
    ['9', 'membership_orders', 'Đơn mua gói (pending/completed/cancelled)'],
    ['10', 'wallet_transactions', 'Giao dịch ví (deposit/purchase/refund/withdraw)'],
    ['11', 'support_tickets', 'Ticket hỗ trợ từ người dùng'],
    ['12', 'activation_codes', 'Mã kích hoạt Pro'],
    ['13', 'bookmarks', 'Từ vựng đã bookmark'],
    ['14', 'user_progress', 'Tiến độ học theo chủ đề'],
    ['15', 'xp_history', 'Lịch sử XP + hoạt động'],
]
for r, data in enumerate(db_data):
    for c, text in enumerate(data):
        db_table.rows[r+1].cells[c].text = text

add_para("", 'Normal')

add_level1("3.2. Xây dựng giao diện ứng dụng")
add_body("Giao diện website English Learning được thiết kế theo phong cách hiện đại, sử dụng tone màu xanh dương (blue) làm chủ đạo, phối hợp với nền sáng (light mode) tạo cảm giác thân thiện và dễ nhìn. Giao diện responsive, tương thích cả PC và mobile.")

# Screenshots for each major page
pages_screenshots = [
    ("3.2.1. Trang đăng nhập", "01_login.png", "Hình 3.9: Giao diện đăng nhập",
     "Trang đăng nhập cho phép người dùng đăng nhập bằng username/password hoặc tài khoản Google. Giao diện đơn giản, rõ ràng với link đăng ký tài khoản mới."),
    
    ("3.2.2. Trang chủ", "02_homepage_top.png", "Hình 3.10: Giao diện trang chủ",
     "Trang chủ hiển thị banner giới thiệu, thống kê nhanh (số chủ đề, bài học, quiz), và danh sách chủ đề học nổi bật. Giao diện được thiết kế hấp dẫn để thu hút người dùng."),

    ("3.2.3. Danh sách chủ đề", "04_topics.png", "Hình 3.11: Giao diện danh sách chủ đề",
     "Hiển thị grid các chủ đề học theo level (Beginner, Intermediate, Advanced). Mỗi chủ đề có thumbnail, tên, mô tả và badge level với màu sắc phân biệt."),
    
    ("3.2.4. Chi tiết chủ đề", "05_topic_detail.png", "Hình 3.12: Giao diện chi tiết chủ đề",
     "Trang chi tiết chủ đề hiển thị danh sách từ vựng (word, phát âm, nghĩa, ví dụ), nút học Flashcard, làm bài test, và danh sách bài học."),

    ("3.2.5. Flashcard", "07_flashcard.png", "Hình 3.13: Giao diện Flashcard",
     "Tính năng Flashcard cho phép người dùng học từ vựng qua thẻ lật (flip card). Mặt trước hiển thị từ tiếng Anh + phiên âm IPA, mặt sau hiển thị nghĩa tiếng Việt + câu ví dụ."),

    ("3.2.6. Bài kiểm tra", "09_test.png", "Hình 3.14: Giao diện danh sách bài kiểm tra",
     "Danh sách bài test phân loại theo: Quiz (miễn phí), Listening và Reading (dành cho Pro). Hiển thị thời gian, số câu, điểm đạt cho mỗi bài."),

    ("3.2.7. Luyện nói (Speaking)", "10_speaking.png", "Hình 3.15: Giao diện luyện nói",
     "Tính năng luyện nói tích hợp Web Speech API để ghi âm giọng nói. AI (GPT) chấm điểm chi tiết: Pronunciation, Fluency, Accuracy và cung cấp feedback cải thiện."),

    ("3.2.8. Dashboard", "11_dashboard.png", "Hình 3.16: Giao diện Dashboard",
     "Dashboard cá nhân hiển thị thống kê tiến độ học tập: streak (chuỗi ngày), XP + Level, mục tiêu hàng ngày, và biểu đồ tiến độ theo chủ đề."),

    ("3.2.9. Ví điện tử", "15_wallet.png", "Hình 3.17: Giao diện ví điện tử",
     "Trang ví hiển thị số dư, nút nạp tiền / rút tiền / mua gói Pro, và lịch sử giao dịch chi tiết với trạng thái (Hoàn tất, Chờ duyệt, Từ chối)."),

    ("3.2.10. Nạp tiền", "16_deposit.png", "Hình 3.18: Giao diện nạp tiền",
     "Trang nạp tiền tích hợp VietQR: người dùng nhập số tiền (có quick buttons 50K-500K), hệ thống tự động tạo QR code chứa thông tin chuyển khoản. Sau khi chuyển khoản, Admin duyệt để cộng balance."),

    ("3.2.11. Nâng cấp Pro", "18_membership_pricing.png", "Hình 3.19: Giao diện nâng cấp Pro",
     "Trang membership hiển thị bảng giá 3 gói (1 tháng, 3 tháng, 12 tháng), so sánh tính năng Free vs Pro. Thanh toán bằng ví với modal xác nhận."),
]

for title, img_file, caption, desc in pages_screenshots:
    add_level2(title)
    add_body(desc)
    img_path = os.path.join(SCREENSHOTS_DIR, img_file)
    add_image(img_path, caption, 5.5)
    add_para("", 'Normal')

# Admin pages
add_level1("3.3. Giao diện quản trị (Admin)")
add_body("Trang quản trị dành cho Admin với đầy đủ chức năng quản lý hệ thống:")

admin_screenshots = [
    ("3.3.1. Dashboard Admin", "22_admin_dashboard.png", "Hình 3.20: Dashboard Admin",
     "Dashboard tổng quan hiển thị: số Users, Pro Members, Chủ đề, Lượt làm bài, Tickets chờ. Kèm theo 4 biểu đồ: tăng trưởng user, phân bố điểm test, đơn nâng cấp, tỷ lệ Free/Pro."),
    ("3.3.2. Quản lý Đơn nâng cấp", "27_admin_orders.png", "Hình 3.21: Quản lý đơn nâng cấp",
     "Danh sách đơn hàng membership. Admin có thể duyệt hoặc từ chối đơn pending."),
    ("3.3.3. Quản lý Ví", "29_admin_wallet.png", "Hình 3.22: Quản lý giao dịch ví",
     "Danh sách giao dịch ví (nạp/rút/mua/hoàn). Admin duyệt/từ chối giao dịch nạp/rút tiền pending."),
    ("3.3.4. Quản lý Tickets", "28_admin_tickets.png", "Hình 3.23: Quản lý Tickets hỗ trợ",
     "Danh sách tickets từ users. Admin phản hồi, duyệt hủy đơn + hoàn tiền tự động vào ví."),
]

for title, img_file, caption, desc in admin_screenshots:
    add_level2(title)
    add_body(desc)
    img_path = os.path.join(SCREENSHOTS_DIR, img_file)
    add_image(img_path, caption, 5.5)
    add_para("", 'Normal')

add_level1("3.4. Xây dựng chức năng")
add_body("Hệ thống English Learning bao gồm các nhóm chức năng chính sau:")

add_level2("3.4.1. Quản lý tài khoản")
add_bullet("Đăng ký tài khoản mới (username, email, password).")
add_bullet("Đăng nhập bằng username/password hoặc Google OAuth 2.0.")
add_bullet("Quản lý hồ sơ: cập nhật thông tin, đổi avatar, đổi mật khẩu.")
add_bullet("Phân quyền: Student (Free/Pro) và Admin.")

add_level2("3.4.2. Hệ thống học tập")
add_bullet("Chủ đề từ vựng: 6 chủ đề (Travel, Food, Business...), mỗi chủ đề có từ vựng, bài học, bài kiểm tra.")
add_bullet("Flashcard: Thẻ lật hiển thị từ vựng, bookmark từ yêu thích.")
add_bullet("Bài học (Lesson): Nội dung đa phương tiện (text, hình, audio, video), đánh giá sao.")
add_bullet("Ngữ pháp: 8 bài ngữ pháp với quiz trắc nghiệm.")
add_bullet("Bài kiểm tra: Quiz (Free), Listening + Reading (Pro), đếm thời gian, chấm điểm tự động.")
add_bullet("Luyện nói (Speaking): Ghi âm → Speech-to-Text → AI đánh giá → Feedback chi tiết.")

add_level2("3.4.3. Hệ thống Gamification")
add_bullet("XP (Experience Points): Nhận XP khi hoàn thành hoạt động học tập.")
add_bullet("Level: Nâng level khi đạt đủ XP (mỗi level cần 100 XP).")
add_bullet("Streak: Đếm số ngày học liên tiếp, tặng bonus XP khi duy trì streak.")
add_bullet("Leaderboard: Bảng xếp hạng top users theo XP.")
add_bullet("Daily Goal: Mục tiêu XP hàng ngày (mặc định 50 XP).")

add_level2("3.4.4. Hệ thống tài chính")
add_bullet("Ví điện tử: Mỗi user có balance riêng, ghi nhận tất cả giao dịch.")
add_bullet("Nạp tiền: User chuyển khoản → Admin duyệt → Cộng balance.")
add_bullet("Mua gói Pro: Trừ balance → Tạo order completed → Nâng cấp Pro.")
add_bullet("Rút tiền: User yêu cầu → Admin chuyển khoản → Trừ balance.")
add_bullet("Hoàn tiền: Khi hủy đơn → Admin duyệt → Hoàn tiền vào ví (tỷ lệ theo thời gian).")
add_bullet("Bảo mật: Sử dụng transaction + row locking (FOR UPDATE) để tránh race condition.")

add_level2("3.4.5. Hệ thống hỗ trợ")
add_bullet("Ticket hỗ trợ: User gửi ticket (Hỗ trợ chung, Hủy đơn, Báo lỗi, Góp ý).")
add_bullet("Admin phản hồi: Trả lời ticket, đổi trạng thái (Open → In Progress → Resolved → Closed).")
add_bullet("Hủy đơn + Hoàn tiền: Chính sách hủy theo thời gian (trong 24h: 100%, 1-7 ngày: 50%, sau 7 ngày: không hoàn).")

add_level1("3.5. Kiểm thử")
add_body("Quá trình kiểm thử được thực hiện trên 30 trang web, bao gồm cả giao diện người dùng và quản trị. Kết quả kiểm thử:")
add_bullet("29/29 trang hoạt động bình thường, không lỗi PHP.")
add_bullet("Vietnamese encoding hiển thị chính xác trên tất cả các trang.")
add_bullet("Wallet + Membership flows hoạt động đúng: nạp tiền, mua gói, hủy đơn, hoàn tiền.")
add_bullet("Google OAuth đăng nhập thành công.")
add_bullet("Speaking AI chấm điểm chính xác.")
add_bullet("Responsive design hoạt động tốt trên PC và mobile.")
add_page_break()

# ============================================================
# CHƯƠNG 4: KẾT LUẬN VÀ KIẾN NGHỊ
# ============================================================
add_heading1("KẾT LUẬN VÀ KIẾN NGHỊ")

add_level1("4.1. Kết quả đạt được")

add_level2("4.1.1. Ưu điểm")
add_bullet("Hoàn thành đầy đủ các chức năng đề ra: học từ vựng, ngữ pháp, bài kiểm tra, luyện nói, flashcard, bookmark.")
add_bullet("Tích hợp thành công AI (OpenAI GPT) vào đánh giá kỹ năng nói – tính năng nổi bật so với các website tương tự.")
add_bullet("Hệ thống ví điện tử + membership chuyên nghiệp với đầy đủ flow: nạp, mua, rút, hoàn tiền.")
add_bullet("Giao diện người dùng đẹp, hiện đại, responsive, trải nghiệm mượt mà.")
add_bullet("Hệ thống gamification (XP, Level, Streak, Leaderboard) tạo động lực học tập.")
add_bullet("Bảo mật giao dịch tài chính: sử dụng database transaction + row locking tránh race condition.")
add_bullet("Code sạch, tổ chức theo mô hình MVC, dễ bảo trì và mở rộng.")
add_bullet("Đăng nhập bằng Google OAuth 2.0 tiện lợi cho người dùng.")

add_level2("4.1.2. Hạn chế")
add_bullet("Tính năng nạp tiền chưa tích hợp tự động (cần Admin duyệt thủ công).")
add_bullet("Chưa có ứng dụng mobile native (chỉ hoạt động trên trình duyệt web).")
add_bullet("Nội dung học tập còn hạn chế (6 chủ đề), cần bổ sung thêm.")
add_bullet("Chưa có tính năng chat/forum cho người dùng trao đổi.")
add_bullet("Speaking chỉ hoạt động trên trình duyệt hỗ trợ Web Speech API (Chrome, Edge).")

add_level1("4.2. Hướng phát triển trong tương lai")
add_bullet("Tích hợp cổng thanh toán tự động: VNPay, Momo, ZaloPay để xác nhận nạp tiền real-time.")
add_bullet("Phát triển ứng dụng mobile bằng React Native hoặc Flutter.")
add_bullet("Bổ sung thêm nội dung: nhiều chủ đề hơn, video bài giảng, podcast tiếng Anh.")
add_bullet("Tích hợp Spaced Repetition System (SRS) để tối ưu hóa việc ôn tập từ vựng.")
add_bullet("Thêm tính năng học nhóm, diễn đàn trao đổi giữa người dùng.")
add_bullet("Sử dụng AI nâng cao hơn: chatbot hội thoại, tự động tạo đề thi, nhận diện giọng nói chính xác hơn.")
add_bullet("Deploy lên hosting thật (VPS, Cloud) để phục vụ người dùng thực tế.")
add_page_break()

# ============================================================
# TÀI LIỆU THAM KHẢO
# ============================================================
add_heading1("TÀI LIỆU THAM KHẢO")

references = [
    "[1]\tNguyễn Văn Hiệp (2023), Giáo trình Lập trình Web với PHP và MySQL, NXB Giáo dục Việt Nam.",
    "[2]\tMatt Zandstra (2021), PHP 8 Objects, Patterns, and Practice, Apress, 6th Edition.",
    "[3]\tLuke Welling, Laura Thomson (2021), PHP and MySQL Web Development, Addison-Wesley, 5th Edition.",
    "[4]\tRobin Nixon (2021), Learning PHP, MySQL & JavaScript, O'Reilly Media, 6th Edition.",
]
add_para("Tiếng Việt và Sách", 'Normal', bold=True)
for ref in references:
    add_para(ref, 'Normal')
add_para("", 'Normal')

web_refs = [
    "[5]\tPHP Official Documentation. Địa chỉ: https://www.php.net/docs.php",
    "[6]\tMySQL Reference Manual. Địa chỉ: https://dev.mysql.com/doc/refman/8.0/en/",
    "[7]\tMDN Web Docs – HTML, CSS, JavaScript. Địa chỉ: https://developer.mozilla.org/",
    "[8]\tGoogle OAuth 2.0 Documentation. Địa chỉ: https://developers.google.com/identity/protocols/oauth2",
    "[9]\tOpenAI API Documentation. Địa chỉ: https://platform.openai.com/docs/",
    "[10]\tWeb Speech API – MDN. Địa chỉ: https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API",
    "[11]\tVietQR API Documentation. Địa chỉ: https://www.vietqr.io/",
    "[12]\tChart.js Documentation. Địa chỉ: https://www.chartjs.org/docs/",
    "[13]\tW3Schools – PHP Tutorial. Địa chỉ: https://www.w3schools.com/php/",
]
add_para("Trang web", 'Normal', bold=True)
for ref in web_refs:
    add_para(ref, 'Normal')

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
print(f"Done! Saved to: {OUTPUT}")
print(f"File size: {os.path.getsize(OUTPUT) / 1024:.0f} KB")
