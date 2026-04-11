"""
Script tạo 2 file .docx cho DATN:
1. docs/wireframe.docx - Thiết kế giao diện (30 screenshots)
2. docs/uml_diagrams.docx - Phân tích hệ thống UML
"""
import os
import base64
import urllib.request
import urllib.parse
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DOCS_DIR = r"g:\xamp\htdocs\DATN\docs"
SCREENSHOTS_DIR = os.path.join(DOCS_DIR, "screenshots")
UML_IMAGES_DIR = os.path.join(DOCS_DIR, "uml_images")
os.makedirs(UML_IMAGES_DIR, exist_ok=True)

def set_style(doc):
    """Configure document styles"""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
        heading_style.font.color.rgb = RGBColor(0, 0, 0)

def add_title_page(doc, title, subtitle=""):
    """Add a title page"""
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = 'Times New Roman'
    
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(14)
        run2.font.name = 'Times New Roman'
        run2.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

def add_screenshot(doc, img_path, caption="", width=6.0):
    """Add screenshot with caption"""
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if caption:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        doc.add_paragraph(f"[Hình ảnh không tìm thấy: {img_path}]")

# ============================================================
# RENDER MERMAID DIAGRAMS AS IMAGES
# ============================================================
def render_mermaid(code, filename):
    """Render Mermaid diagram to PNG using mermaid.ink"""
    filepath = os.path.join(UML_IMAGES_DIR, filename)
    if os.path.exists(filepath) and os.path.getsize(filepath) > 1000:
        print(f"  Cached: {filename}")
        return filepath
    
    try:
        encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}?bgColor=white&theme=default"
        
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=30) as response:
            data = response.read()
        
        with open(filepath, 'wb') as f:
            f.write(data)
        print(f"  Rendered: {filename} ({len(data)} bytes)")
        return filepath
    except Exception as e:
        print(f"  ERROR rendering {filename}: {e}")
        return None

# Mermaid diagram codes
diagrams = {
    "use_case.png": """graph TB
    subgraph "Hệ thống EnglishMaster"
        UC1["Đăng ký / Đăng nhập"]
        UC2["Đăng nhập Google OAuth"]
        UC3["Xem chủ đề từ vựng"]
        UC4["Học bài Lesson"]
        UC5["Học Flashcard"]
        UC6["Làm bài test Quiz"]
        UC7["Làm bài Listening"]
        UC8["Làm bài Reading"]
        UC9["Luyện nói Speaking"]
        UC10["Học Ngữ pháp"]
        UC11["Đánh giá bài học"]
        UC12["Bookmark từ vựng"]
        UC13["Tìm kiếm"]
        UC14["Xem Dashboard"]
        UC15["Xem Leaderboard"]
        UC16["Nâng cấp Pro"]
        UC17["Nạp tiền vào ví"]
        UC18["Mua gói bằng ví"]
        UC19["Rút tiền"]
        UC20["Gửi ticket hỗ trợ"]
        UC21["Hủy đơn"]
        UC22["Quản lý Users"]
        UC23["Quản lý Chủ đề"]
        UC24["Duyệt đơn"]
        UC25["Duyệt giao dịch ví"]
        UC26["Quản lý Tickets"]
    end
    Guest(("Khách"))
    User(("User"))
    Pro(("Pro User"))
    Admin(("Admin"))
    Guest --> UC1
    Guest --> UC2
    Guest --> UC3
    User --> UC4
    User --> UC5
    User --> UC6
    User --> UC10
    User --> UC11
    User --> UC12
    User --> UC13
    User --> UC14
    User --> UC15
    User --> UC16
    User --> UC17
    User --> UC18
    User --> UC19
    User --> UC20
    User --> UC21
    Pro --> UC7
    Pro --> UC8
    Pro --> UC9
    Admin --> UC22
    Admin --> UC23
    Admin --> UC24
    Admin --> UC25
    Admin --> UC26""",

    "sequence_login.png": """sequenceDiagram
    actor U as User
    participant V as Login Page
    participant C as AuthController
    participant M as User Model
    participant DB as Database
    participant S as Session
    U->>V: Nhập username + password
    V->>C: POST /auth/login
    C->>M: authenticate(username, password)
    M->>DB: SELECT * FROM users WHERE username = ?
    DB-->>M: User data
    M->>M: password_verify(password, hash)
    alt Thành công
        M-->>C: User object
        C->>S: Lưu session
        C-->>V: Redirect Trang chủ
    else Thất bại
        M-->>C: null
        C-->>V: Flash error
    end""",

    "sequence_deposit.png": """sequenceDiagram
    actor U as User
    participant V as Deposit Page
    participant W as WalletController
    participant DB as Database
    actor A as Admin
    participant AC as AdminController
    U->>V: Nhập số tiền 50,000đ
    V->>V: Hiển thị QR VietQR
    U->>U: Chuyển khoản
    U->>V: Click Đã chuyển khoản
    V->>W: POST /wallet/createDeposit
    W->>DB: INSERT wallet_transactions status=pending
    W-->>V: Chờ Admin duyệt
    Note over A,AC: Admin xác nhận
    A->>AC: POST /admin/approveTransaction
    AC->>DB: BEGIN TRANSACTION
    AC->>DB: UPDATE users balance += amount
    AC->>DB: UPDATE wallet_transactions status=completed
    AC->>DB: COMMIT
    AC-->>A: Duyệt thành công""",

    "sequence_purchase.png": """sequenceDiagram
    actor U as User
    participant V as Membership Page
    participant MC as MembershipController
    participant DB as Database
    U->>V: Click Chọn gói Pro 1 Tháng
    V->>V: Hiện modal thanh toán
    U->>V: Click Thanh toán bằng ví
    V->>MC: POST /membership/createOrder
    MC->>DB: BEGIN TRANSACTION
    MC->>DB: SELECT balance FOR UPDATE
    alt Đủ tiền
        MC->>DB: UPDATE users balance -= price
        MC->>DB: INSERT membership_orders completed
        MC->>DB: INSERT wallet_transactions purchase
        MC->>DB: COMMIT
        MC-->>V: Mua thành công!
    else Không đủ tiền
        MC->>DB: ROLLBACK
        MC-->>V: Số dư không đủ
    end""",

    "sequence_cancel.png": """sequenceDiagram
    actor U as User
    participant V as Support Page
    participant SC as SupportController
    participant DB as Database
    actor A as Admin
    participant AC as AdminController
    U->>V: Chọn Hủy đơn + chọn order
    V->>SC: POST /support/store
    SC->>SC: checkCancelEligibility
    SC->>DB: INSERT support_tickets
    SC-->>V: Đã gửi yêu cầu hủy
    Note over A,AC: Admin duyệt hủy
    A->>AC: POST /admin/approveCancelOrder
    AC->>DB: BEGIN TRANSACTION
    AC->>DB: UPDATE orders cancelled
    AC->>DB: UPDATE users balance += refund
    AC->>DB: INSERT wallet_transactions refund
    AC->>DB: UPDATE ticket resolved
    AC->>DB: COMMIT
    AC-->>A: Đã hủy + hoàn tiền""",

    "state_order.png": """stateDiagram-v2
    [*] --> Pending: User chuyển khoản
    [*] --> Completed: Mua bằng ví hoặc mã kích hoạt
    Pending --> Completed: Admin duyệt
    Pending --> Cancelled: Admin từ chối
    Pending --> Cancelled: User hủy đơn
    Completed --> [*]
    Cancelled --> [*]""",

    "state_wallet.png": """stateDiagram-v2
    [*] --> Pending: User tạo yêu cầu nạp hoặc rút
    [*] --> Completed: Hệ thống tự ghi purchase hoặc refund
    Pending --> Completed: Admin duyệt
    Pending --> Rejected: Admin từ chối
    Completed --> [*]
    Rejected --> [*]""",

    "state_ticket.png": """stateDiagram-v2
    [*] --> Open: User gửi ticket
    Open --> InProgress: Admin đang xử lý
    Open --> Resolved: Admin phản hồi
    InProgress --> Resolved: Admin phản hồi
    Resolved --> Closed: Admin đóng
    Closed --> [*]""",

    "state_membership.png": """stateDiagram-v2
    [*] --> Free: Đăng ký tài khoản
    Free --> Pro: Mua gói hoặc kích hoạt mã
    Pro --> Free: Hết hạn membership
    Pro --> Pro: Gia hạn thêm gói"""
}

print("=== Rendering Mermaid Diagrams ===")
rendered = {}
for filename, code in diagrams.items():
    result = render_mermaid(code, filename)
    rendered[filename] = result

# ============================================================
# CREATE WIREFRAME.DOCX
# ============================================================
print("\n=== Creating wireframe.docx ===")
doc = Document()
set_style(doc)
add_title_page(doc, "THIẾT KẾ GIAO DIỆN", "Hệ thống học tiếng Anh trực tuyến — EnglishMaster")

# Section 1: User pages
doc.add_heading("I. GIAO DIỆN NGƯỜI DÙNG (STUDENT)", level=1)

pages = [
    ("1. Trang đăng nhập", "01_login.png", [
        "Form đăng nhập username/password",
        "Nút đăng nhập bằng Google OAuth 2.0",
        "Link đăng ký tài khoản mới"
    ]),
    ("2. Trang chủ", ["02_homepage_top.png", "03_homepage_bottom.png"], [
        "Banner giới thiệu hệ thống",
        "Thống kê nhanh (số chủ đề, bài học, quiz)",
        "Danh sách chủ đề học nổi bật",
        "Footer với links hữu ích"
    ]),
    ("3. Danh sách chủ đề", "04_topics.png", [
        "Grid các chủ đề theo level (Beginner, Intermediate, Advanced)",
        "Hiển thị thumbnail, tên, mô tả ngắn",
        "Badge level với màu sắc phân biệt"
    ]),
    ("4. Chi tiết chủ đề", "05_topic_detail.png", [
        "Thông tin chủ đề (tên, level, mô tả)",
        "Danh sách từ vựng (word, nghĩa, phát âm, ví dụ)",
        "Nút học Flashcard, làm bài test",
        "Danh sách bài học trong chủ đề"
    ]),
    ("5. Bài học (Lesson)", "06_lesson.png", [
        "Nội dung bài học đa phương tiện (text, hình ảnh, audio, video)",
        "Breadcrumb navigation",
        "Nút Next/Previous lesson",
        "Phần đánh giá bài học (rating 1-5 sao + nhận xét)"
    ]),
    ("6. Flashcard", "07_flashcard.png", [
        "Card lật (flip) hiển thị từ vựng",
        "Mặt trước: từ tiếng Anh + phiên âm IPA",
        "Mặt sau: nghĩa tiếng Việt + câu ví dụ",
        "Nút Previous/Next, nút Bookmark lưu từ"
    ]),
    ("7. Ngữ pháp (Grammar)", "08_grammar.png", [
        "Phân loại theo chủ đề: Tenses, Prepositions, Modals...",
        "Badge level cho từng bài (Beginner, Intermediate, Advanced)",
        "Hiển thị số câu quiz mỗi bài"
    ]),
    ("8. Bài kiểm tra (Test)", "09_test.png", [
        "Danh sách bài test theo chủ đề",
        "Phân loại: Quiz, Listening (Pro), Reading (Pro)",
        "Hiển thị thời gian, số câu, điểm đạt",
        "Badge PRO cho tính năng nâng cao"
    ]),
    ("9. Luyện nói (Speaking)", "10_speaking.png", [
        "Chọn chủ đề luyện nói",
        "Ghi âm giọng nói trực tiếp trên trình duyệt",
        "AI chấm điểm: Pronunciation, Fluency, Accuracy",
        "Feedback chi tiết từ AI"
    ]),
    ("10. Dashboard", "11_dashboard.png", [
        "Thống kê tiến độ học tập cá nhân",
        "Streak (chuỗi ngày học liên tiếp)",
        "XP và Level hiện tại",
        "Biểu đồ tiến độ theo từng chủ đề"
    ]),
    ("11. Bảng xếp hạng", "12_leaderboard.png", [
        "Top users theo XP",
        "Hiển thị avatar, tên, level, streak",
        "Highlight vị trí hiện tại của user"
    ]),
    ("12. Bookmark", "13_bookmark.png", [
        "Danh sách từ vựng đã lưu (bookmark)",
        "Nút xóa bookmark",
        "Ghi chú cá nhân cho mỗi từ"
    ]),
    ("13. Hồ sơ cá nhân", "14_profile.png", [
        "Thông tin tài khoản (username, email, tên)",
        "Upload/đổi avatar",
        "Đổi mật khẩu",
        "Trạng thái membership (Free/Pro)"
    ]),
    ("14. Ví điện tử", "15_wallet.png", [
        "Hiển thị số dư khả dụng",
        "Nút Nạp tiền, Rút tiền, Mua gói Pro",
        "Lịch sử giao dịch (bảng chi tiết với trạng thái)"
    ]),
    ("15. Nạp tiền", "16_deposit.png", [
        "Nhập số tiền (quick buttons: 50K, 100K, 200K, 500K)",
        "Hiển thị QR code VietQR tự động",
        "Thông tin chuyển khoản (STK, chủ TK, nội dung CK)",
        "Nút xác nhận đã chuyển khoản"
    ]),
    ("16. Rút tiền", "17_withdraw.png", [
        "Form nhập số tiền rút (tối thiểu 50,000đ)",
        "Chọn ngân hàng (VCB, TCB, BIDV, MBBank...)",
        "Nhập STK + tên chủ tài khoản",
        "Thông báo thời gian xử lý 1-3 ngày làm việc"
    ]),
    ("17. Nâng cấp Pro", ["18_membership_pricing.png", "19_membership_history.png"], [
        "So sánh tính năng Free vs Pro",
        "Bảng giá 3 gói: 1 tháng, 3 tháng, 12 tháng",
        "Badge 'PHỔ BIẾN NHẤT' cho gói 3 tháng",
        "Modal thanh toán bằng ví",
        "Lịch sử đơn hàng + nút hủy đơn"
    ]),
    ("18. Hỗ trợ (Support)", ["20_support_list.png", "21_support_create.png"], [
        "Danh sách ticket đã gửi (trạng thái: Mở, Đang xử lý, Đã xử lý)",
        "Form tạo ticket mới (Hỗ trợ chung, Hủy đơn, Báo lỗi, Góp ý)",
        "Hiển thị phản hồi từ Admin"
    ]),
]

for title, imgs, features in pages:
    doc.add_heading(title, level=2)
    
    # Add screenshots
    if isinstance(imgs, list):
        for img in imgs:
            add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, img), width=5.8)
            doc.add_paragraph("")
    else:
        add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, imgs), width=5.8)
        doc.add_paragraph("")
    
    # Add feature list
    p = doc.add_paragraph()
    run = p.add_run("Chức năng chính:")
    run.bold = True
    run.font.size = Pt(13)
    
    for feature in features:
        p = doc.add_paragraph(feature, style='List Bullet')
        p.style.font.size = Pt(12)
    
    doc.add_paragraph("")

# Section 2: Admin pages
doc.add_page_break()
doc.add_heading("II. GIAO DIỆN QUẢN TRỊ (ADMIN)", level=1)

admin_pages = [
    ("19. Dashboard Admin", "22_admin_dashboard.png", [
        "Thống kê tổng quan: Users, Pro Members, Chủ đề, Lượt làm bài, Tickets",
        "Biểu đồ tăng trưởng User (7 ngày gần nhất)",
        "Biểu đồ phân bố điểm test",
        "Biểu đồ đơn nâng cấp theo tháng (6 tháng)",
        "Biểu đồ tỷ lệ Free/Pro (Pie chart)"
    ]),
    ("20. Quản lý Users", "23_admin_users.png", [
        "Danh sách tất cả users (phân trang)",
        "Tìm kiếm user theo username/email",
        "Sửa thông tin, đổi role (student/admin)",
        "Ban/unban tài khoản"
    ]),
    ("21. Quản lý Chủ đề", "24_admin_topics.png", [
        "CRUD chủ đề học (thêm, sửa, xóa)",
        "Sắp xếp thứ tự hiển thị",
        "Bật/tắt active"
    ]),
    ("22. Quản lý Câu hỏi", "25_admin_questions.png", [
        "Danh sách câu hỏi theo bài test",
        "CRUD câu hỏi (multiple choice, true/false, fill blank)",
        "Filter theo loại test"
    ]),
    ("23. Quản lý Mã kích hoạt", "26_admin_codes.png", [
        "Tạo mã kích hoạt Pro tự động",
        "Danh sách mã (đã dùng / chưa dùng)",
        "Thông tin user đã sử dụng mã"
    ]),
    ("24. Quản lý Đơn nâng cấp", "27_admin_orders.png", [
        "Danh sách đơn hàng membership",
        "Duyệt / Từ chối đơn pending",
        "Xem thông tin chuyển khoản"
    ]),
    ("25. Quản lý Tickets", "28_admin_tickets.png", [
        "Danh sách tất cả tickets từ users",
        "Phản hồi ticket trực tiếp",
        "Duyệt hủy đơn + hoàn tiền tự động vào ví",
        "Thay đổi trạng thái ticket"
    ]),
    ("26. Quản lý Ví", "29_admin_wallet.png", [
        "Danh sách giao dịch ví (nạp / rút / mua / hoàn tiền)",
        "Duyệt / Từ chối giao dịch pending",
        "Hiển thị thông tin ngân hàng (cho lệnh rút)"
    ]),
    ("27. Cài đặt hệ thống", "30_admin_settings.png", [
        "Cấu hình các thông số hệ thống",
        "Quản lý cài đặt chung"
    ]),
]

for title, img, features in admin_pages:
    doc.add_heading(title, level=2)
    add_screenshot(doc, os.path.join(SCREENSHOTS_DIR, img), width=5.8)
    doc.add_paragraph("")
    
    p = doc.add_paragraph()
    run = p.add_run("Chức năng chính:")
    run.bold = True
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    doc.add_paragraph("")

wireframe_path = os.path.join(DOCS_DIR, "wireframe.docx")
doc.save(wireframe_path)
print(f"Saved: {wireframe_path}")

# ============================================================
# CREATE UML_DIAGRAMS.DOCX
# ============================================================
print("\n=== Creating uml_diagrams.docx ===")
doc2 = Document()
set_style(doc2)
add_title_page(doc2, "PHÂN TÍCH HỆ THỐNG", "Sơ đồ UML — EnglishMaster")

# 1. Use Case Diagram
doc2.add_heading("1. Use Case Diagram", level=1)
doc2.add_paragraph("Sơ đồ Use Case mô tả các chức năng chính của hệ thống và mối quan hệ giữa các actor (tác nhân) với hệ thống.")

doc2.add_heading("1.1. Danh sách Actor", level=2)
table = doc2.add_table(rows=5, cols=3)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Actor", "Mô tả", "Quyền hạn"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True

actors = [
    ["Guest (Khách)", "Người dùng chưa đăng nhập", "Xem chủ đề, đăng ký, đăng nhập"],
    ["User (Học viên)", "Người dùng đã đăng nhập (Free)", "Học từ vựng, làm quiz, flashcard, ngữ pháp, bookmark, ví điện tử"],
    ["Pro User", "Hội viên Pro (trả phí)", "Tất cả quyền User + Listening, Reading, Speaking"],
    ["Admin", "Quản trị viên", "Tất cả quyền Pro + Quản lý hệ thống, duyệt đơn, quản lý ví"]
]
for r, row_data in enumerate(actors):
    for c, text in enumerate(row_data):
        table.rows[r+1].cells[c].text = text

doc2.add_paragraph("")
doc2.add_heading("1.2. Sơ đồ Use Case", level=2)

if rendered.get("use_case.png"):
    add_screenshot(doc2, rendered["use_case.png"], "Hình 1: Sơ đồ Use Case tổng quan", 5.5)
else:
    doc2.add_paragraph("[Sơ đồ Use Case - xem file uml_diagrams.md]")

doc2.add_paragraph("")

# Use Case Table
doc2.add_heading("1.3. Mô tả Use Case chính", level=2)
uc_table = doc2.add_table(rows=13, cols=4)
uc_table.style = 'Table Grid'
uc_headers = ["STT", "Use Case", "Actor", "Mô tả"]
for i, h in enumerate(uc_headers):
    uc_table.rows[0].cells[i].text = h
    for p in uc_table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

uc_data = [
    ["1", "Đăng ký / Đăng nhập", "Guest", "Tạo tài khoản hoặc đăng nhập bằng username/password hoặc Google OAuth"],
    ["2", "Học từ vựng (Flashcard)", "User", "Xem và học từ vựng theo chủ đề với thẻ lật (flip card)"],
    ["3", "Làm bài test Quiz", "User", "Làm bài kiểm tra trắc nghiệm, điền từ, đúng/sai"],
    ["4", "Luyện nói (Speaking)", "Pro", "Ghi âm giọng nói, AI chấm điểm phát âm + trôi chảy"],
    ["5", "Làm bài Listening", "Pro", "Nghe audio và trả lời câu hỏi"],
    ["6", "Nạp tiền vào ví", "User", "Chuyển khoản ngân hàng → Admin duyệt → Cộng balance"],
    ["7", "Mua gói Pro", "User", "Thanh toán bằng ví hoặc mã kích hoạt để nâng cấp Pro"],
    ["8", "Hủy đơn + Hoàn tiền", "User", "Gửi yêu cầu hủy → Admin duyệt → Hoàn tiền vào ví"],
    ["9", "Rút tiền", "User", "Gửi yêu cầu rút → Admin chuyển khoản → Trừ balance"],
    ["10", "Duyệt giao dịch ví", "Admin", "Xác nhận nạp/rút tiền cho user"],
    ["11", "Quản lý nội dung", "Admin", "CRUD chủ đề, bài học, câu hỏi, mã kích hoạt"],
    ["12", "Quản lý tickets", "Admin", "Phản hồi ticket hỗ trợ, duyệt hủy đơn"],
]
for r, row_data in enumerate(uc_data):
    for c, text in enumerate(row_data):
        uc_table.rows[r+1].cells[c].text = text

# 2. Class Diagram
doc2.add_page_break()
doc2.add_heading("2. Class Diagram", level=1)
doc2.add_paragraph("Sơ đồ Class mô tả cấu trúc các lớp (entity) trong hệ thống, thuộc tính và quan hệ giữa chúng.")

# Class table
doc2.add_heading("2.1. Danh sách Entity", level=2)
class_table = doc2.add_table(rows=12, cols=3)
class_table.style = 'Table Grid'
class_headers = ["Entity", "Table", "Mô tả"]
for i, h in enumerate(class_headers):
    class_table.rows[0].cells[i].text = h
    for p in class_table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

classes = [
    ["User", "users", "Người dùng hệ thống (username, email, balance, membership...)"],
    ["Topic", "topics", "Chủ đề học (Travel, Food, Business...)"],
    ["Vocabulary", "vocabularies", "Từ vựng (word, pronunciation, meaning, example)"],
    ["Lesson", "lessons", "Bài học trong chủ đề"],
    ["Test", "tests", "Bài kiểm tra (quiz, listening, reading)"],
    ["Question", "questions", "Câu hỏi trong bài test"],
    ["MembershipPlan", "membership_plans", "Gói Pro (1 tháng, 3 tháng, 12 tháng)"],
    ["MembershipOrder", "membership_orders", "Đơn mua gói Pro"],
    ["WalletTransaction", "wallet_transactions", "Giao dịch ví (nạp, mua, hoàn, rút)"],
    ["SupportTicket", "support_tickets", "Ticket hỗ trợ từ user"],
    ["GrammarLesson", "grammar_lessons", "Bài ngữ pháp English"]
]
for r, row_data in enumerate(classes):
    for c, text in enumerate(row_data):
        class_table.rows[r+1].cells[c].text = text

doc2.add_paragraph("")
doc2.add_heading("2.2. Quan hệ giữa các Entity", level=2)
relations = [
    "User (1) → (*) MembershipOrder: Một user có nhiều đơn hàng",
    "User (1) → (*) WalletTransaction: Một user có nhiều giao dịch ví",
    "User (1) → (*) SupportTicket: Một user có nhiều ticket hỗ trợ",
    "User (1) → (*) UserProgress: Tiến độ theo từng chủ đề",
    "Topic (1) → (*) Vocabulary: Một chủ đề có nhiều từ vựng",
    "Topic (1) → (*) Lesson: Một chủ đề có nhiều bài học",
    "Topic (1) → (*) Test: Một chủ đề có nhiều bài kiểm tra",
    "Test (1) → (*) Question: Một bài test có nhiều câu hỏi",
    "MembershipPlan (1) → (*) MembershipOrder: Một gói có nhiều đơn đặt",
    "MembershipOrder (1) → (0..1) SupportTicket: Đơn hàng liên kết ticket hủy"
]
for rel in relations:
    doc2.add_paragraph(rel, style='List Bullet')

# 3. Sequence Diagrams
doc2.add_page_break()
doc2.add_heading("3. Sequence Diagram", level=1)
doc2.add_paragraph("Sơ đồ Sequence mô tả luồng tương tác giữa các thành phần trong hệ thống theo thời gian.")

seq_diagrams = [
    ("3.1. Luồng đăng nhập", "sequence_login.png", "Hình 2: Sequence Diagram — Đăng nhập"),
    ("3.2. Luồng nạp tiền vào ví", "sequence_deposit.png", "Hình 3: Sequence Diagram — Nạp tiền"),
    ("3.3. Luồng mua gói Pro bằng ví", "sequence_purchase.png", "Hình 4: Sequence Diagram — Mua gói Pro"),
    ("3.4. Luồng hủy đơn + hoàn tiền", "sequence_cancel.png", "Hình 5: Sequence Diagram — Hủy đơn + Hoàn tiền"),
]

for title, img_key, caption in seq_diagrams:
    doc2.add_heading(title, level=2)
    if rendered.get(img_key):
        add_screenshot(doc2, rendered[img_key], caption, 5.8)
    else:
        doc2.add_paragraph(f"[{caption} - xem file uml_diagrams.md]")
    doc2.add_paragraph("")

# 4. State Diagrams
doc2.add_page_break()
doc2.add_heading("4. State Diagram", level=1)
doc2.add_paragraph("Sơ đồ State mô tả các trạng thái có thể có của một đối tượng và sự chuyển đổi giữa chúng.")

state_diagrams = [
    ("4.1. Trạng thái Đơn hàng (Membership Order)", "state_order.png", "Hình 6: State Diagram — Membership Order",
     [("Pending", "Đơn chờ xác nhận thanh toán"), ("Completed", "Đã kích hoạt Pro thành công"), ("Cancelled", "Đã hủy + hoàn tiền (nếu đủ điều kiện)")]),
    ("4.2. Trạng thái Giao dịch Ví (Wallet Transaction)", "state_wallet.png", "Hình 7: State Diagram — Wallet Transaction",
     [("Pending", "Chờ Admin xác nhận"), ("Completed", "Đã xử lý + cập nhật balance"), ("Rejected", "Từ chối + ghi lý do")]),
    ("4.3. Trạng thái Support Ticket", "state_ticket.png", "Hình 8: State Diagram — Support Ticket",
     [("Open", "Mới gửi, chờ xử lý"), ("In Progress", "Admin đang xử lý"), ("Resolved", "Đã giải quyết"), ("Closed", "Đã đóng")]),
    ("4.4. Trạng thái Membership User", "state_membership.png", "Hình 9: State Diagram — User Membership",
     [("Free", "Tài khoản miễn phí, giới hạn tính năng"), ("Pro", "Hội viên Pro, mở khóa tất cả tính năng")]),
]

for title, img_key, caption, states in state_diagrams:
    doc2.add_heading(title, level=2)
    if rendered.get(img_key):
        add_screenshot(doc2, rendered[img_key], caption, 4.5)
    else:
        doc2.add_paragraph(f"[{caption} - xem file uml_diagrams.md]")
    
    doc2.add_paragraph("")
    p = doc2.add_paragraph()
    run = p.add_run("Mô tả trạng thái:")
    run.bold = True
    
    state_table = doc2.add_table(rows=len(states)+1, cols=2)
    state_table.style = 'Table Grid'
    state_table.rows[0].cells[0].text = "Trạng thái"
    state_table.rows[0].cells[1].text = "Mô tả"
    for p in state_table.rows[0].cells[0].paragraphs:
        p.runs[0].bold = True
    for p in state_table.rows[0].cells[1].paragraphs:
        p.runs[0].bold = True
    
    for i, (state, desc) in enumerate(states):
        state_table.rows[i+1].cells[0].text = state
        state_table.rows[i+1].cells[1].text = desc
    
    doc2.add_paragraph("")

uml_path = os.path.join(DOCS_DIR, "uml_diagrams.docx")
doc2.save(uml_path)
print(f"Saved: {uml_path}")

print("\n=== DONE ===")
print(f"1. {wireframe_path}")
print(f"2. {uml_path}")
